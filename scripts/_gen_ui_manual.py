"""
生成 UI自动化说明及配置手册.docx
完全参考项目中现有的 API接口自动化说明及配置手册.docx 风格（推测），
结构：1. 项目简介 2. 环境要求 3. 目录结构 4. 配置说明 5. Pipeline 流程 6. 使用指南 7. 产物说明 8. CI/CD 9. 常见问题
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL

OUT = Path(__file__).resolve().parent.parent / "UI自动化说明及配置手册.docx"

doc = Document()

# ---- 全局样式 ----
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.font.size = Pt(10.5)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    else:
        hs.font.size = Pt(12)


def add_code_block(text: str, indent: int = 28):
    """代码块：浅灰底 + 等宽字体"""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent / 72)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_note(text: str, is_warn: bool = False):
    """注意/警告段落"""
    p = doc.add_paragraph()
    run = p.add_run(("⚠️  " if is_warn else "💡  ") + text)
    run.font.color.rgb = (
        RGBColor(0xC0, 0x00, 0x00) if is_warn else RGBColor(0x1F, 0x4E, 0x79)
    )
    run.font.bold = True


def add_simple_table(headers: list[str], rows: list[list]):
    """通用表格（蓝色表头 + 边框）"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    # 表头
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 蓝底
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): "2E75B6",
            qn("w:val"): "clear",
        }, nsmap=None)
        tcPr.append(shd)
    # 内容
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = str(val)
    return t


# ============================================================
# 标题页
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("UI 自动化\n说明及配置手册")
tr.font.name = "微软雅黑"
tr.font.size = Pt(36)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Python + Playwright + Pytest + Allure + POM + AI 引擎")
sr.font.name = "微软雅黑"
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
vr = ver_p.add_run("文档版本: V1.0    日期: 2026-08-21")
vr.font.name = "微软雅黑"
vr.font.size = Pt(10.5)
vr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
doc.add_heading("目  录", level=1)
toc_entries = [
    ("一、项目简介", "3"),
    ("    1.1 项目定位", "3"),
    ("    1.2 架构总览", "3"),
    ("二、环境要求", "4"),
    ("    2.1 软件版本", "4"),
    ("    2.2 依赖安装", "4"),
    ("三、目录结构", "5"),
    ("四、配置说明", "7"),
    ("    4.1 基础配置 (config.yaml)", "7"),
    ("    4.2 钉钉通知 (notification.json)", "9"),
    ("    4.3 Pytest 配置 (pytest.ini)", "10"),
    ("五、Pipeline 运行流程", "11"),
    ("    5.1 八阶段全链路", "11"),
    ("    5.2 AI 引擎详解", "12"),
    ("    5.3 输入端解析", "13"),
    ("    5.4 自愈机制", "14"),
    ("六、使用指南", "15"),
    ("    6.1 快速开始", "15"),
    ("    6.2 全链路运行", "16"),
    ("    6.3 单模块/筛选运行", "17"),
    ("    6.4 AI 模式说明", "18"),
    ("七、产物说明", "19"),
    ("    7.1 Allure 报告", "19"),
    ("    7.2 Excel 结果", "20"),
    ("    7.3 Bug 清单", "20"),
    ("八、CI/CD 集成", "21"),
    ("九、常见问题 FAQ", "23"),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1 * (len(entry) - len(entry.lstrip())))
    p.add_run(entry.ljust(50, " ")).font.name = "微软雅黑"

doc.add_page_break()

# ============================================================
# 一、项目简介
# ============================================================
doc.add_heading("一、项目简介", level=1)

doc.add_heading("1.1 项目定位", level=2)
doc.add_paragraph(
    "本框架是一套 AI 驱动的 UI 自动化测试解决方案，基于 Python + Playwright + Pytest + Allure + POM 设计模式，"
    "对标 Midscene / zeroStep / AutoPlaywright 等主流 AI 自动化工具。"
)
doc.add_paragraph("核心特性：")
items = [
    "AI 生成用例：支持 PRD 文档、蓝湖设计稿、原型截图三种输入端，AI 自动生成 Page Object + pytest 用例",
    "POM 页面对象模式：所有页面元素与操作封装到 pages/ 下的 Page 类，逻辑清晰可维护",
    "自愈定位器：元素定位失败时自动调用 AI 引擎寻找替代定位，无需手动修改脚本",
    "八阶段 Pipeline：完全对齐 API 自动化框架的执行链路，团队上手零认知成本",
    "Allure 可视化报告：支持步骤级截图 + ECharts 图表 + 失败堆栈 + 视频/ trace 回放",
    "Excel / 钉钉 / CI/CD 闭环：结果自动导出、推送、GitHub Actions 一键触发",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("1.2 架构总览", level=2)
doc.add_paragraph("Pipeline 采用四层分离架构，数据流从左到右单向流转：")
add_code_block(" 输入端(parsers)  →  AI 引擎(ai)  →  制品(pages/tests)  →  执行(utils+playwright)  →  产物(reports/outputs)")
doc.add_paragraph()
add_simple_table(
    ["层级", "目录", "职责"],
    [
        ["输入端", "UI_input_files/ + parsers/", "PRD / 蓝湖 / 原型截图解析为统一 UIRequirement 结构"],
        ["AI 引擎", "ai/", "LLM 调用 + 自愈定位 + 代码生成器 + Prompt 模板库"],
        ["制品层", "pages/ + tests/", "Page Object 页面对象 + pytest 测试用例（AI 生成或手写）"],
        ["执行层", "utils/ + conftest.py", "Playwright 浏览器、断言、截图、Excel 导出、钉钉推送"],
        ["Pipeline", "scripts/run_pipeline.py", "8 阶段一键编排"],
        ["产物层", "reports/ + UI_output_files/", "Allure 报告、Excel 结果、Bug 清单"],
    ],
)

doc.add_page_break()

# ============================================================
# 二、环境要求
# ============================================================
doc.add_heading("二、环境要求", level=1)

doc.add_heading("2.1 软件版本", level=2)
add_simple_table(
    ["软件", "推荐版本", "说明"],
    [
        ["Python", "3.11+（已在 3.11、3.14 验证）", "x64 版本，推荐 Anaconda + venv"],
        ["JDK", "11+", "Allure 是 Java 程序，需配置 JAVA_HOME"],
        ["Allure", "2.30+", "Windows 推荐从 allure-2.34.1.zip 解压后加入 PATH"],
        ["Node.js", "18+（可选）", "npm 方式安装 allure 时需要"],
        ["Chromium 浏览器", "Playwright 管理", "执行 playwright install chromium 自动下载"],
        ["操作系统", "Windows 10+ / macOS 11+ / Linux", "已在 Windows 11 完整验证"],
        ["Git", "2.30+", "Clone 仓库与提交产物"],
    ],
)

doc.add_heading("2.2 依赖安装", level=2)
doc.add_paragraph("按如下顺序安装：")
add_code_block(
    "# 1. Clone 项目\n"
    "git clone https://github.com/ChangXia-TestTeam/TestUi-pub-repository.git\n"
    "cd TestUi-pub-repository\n\n"
    "# 2. 创建并激活虚拟环境（推荐）\n"
    "python -m venv .venv\n"
    ".venv\\Scripts\\activate   # Windows\n"
    "# source .venv/bin/activate  # macOS/Linux\n\n"
    "# 3. 安装 Python 依赖\n"
    "pip install -r requirements.txt\n\n"
    "# 4. 安装 Playwright 浏览器\n"
    "python -m playwright install chromium --with-deps\n\n"
    "# 5. 验证 Allure\n"
    "allure --version   # 应输出 2.34.x"
)
add_note("如果 allure --version 提示找不到命令，手动将 allure-2.34.1/bin 目录加入系统 PATH 变量。")

doc.add_page_break()

# ============================================================
# 三、目录结构
# ============================================================
doc.add_heading("三、目录结构", level=1)
doc.add_paragraph("标准项目树如下，所有新增文件请放入对应目录：")
add_code_block(
    "test_ui_sjpt/\n"
    "├── config/                           # 📁 配置文件（必须修改！）\n"
    "│   ├── config.yaml                  # 环境/账号/浏览器/AI 引擎\n"
    "│   └── notification.json             # 钉钉机器人 Webhook + Secret\n"
    "│\n"
    "├── ai/                               # 🤖 AI 引擎层（核心）\n"
    "│   ├── llm_client.py                 # LLM 调用封装（trae / api 双模式）\n"
    "│   ├── self_healing.py               # 定位失败自愈（Midscene 风格）\n"
    "│   ├── generators/\n"
    "│   │   ├── page_object_gen.py        # 需求 → Page Object 代码\n"
    "│   │   ├── test_case_gen.py          # 需求 → pytest 用例代码\n"
    "│   │   └── locator_gen.py            # 截图 → Playwright 定位\n"
    "│   ├── prompts/                      # Prompt 模板库（可按项目定制）\n"
    "│   │   ├── page_object.md\n"
    "│   │   ├── test_case.md\n"
    "│   │   └── locator.md\n"
    "│   ├── .pending/                     # trae 模式：待执行 AI 任务\n"
    "│   └── .done/                        # trae 模式：已完成 AI 任务\n"
    "│\n"
    "├── parsers/                          # 📑 输入端解析层\n"
    "│   ├── requirement_model.py          # 统一 UIRequirement 数据模型\n"
    "│   ├── prd_parser.py                 # PRD (docx/pdf/md)\n"
    "│   ├── lanhu_parser.py               # 蓝湖导出 / 在线链接\n"
    "│   └── screenshot_parser.py          # 原型截图（AI 补充定位）\n"
    "│\n"
    "├── pages/                            # 🧩 POM 页面对象层\n"
    "│   ├── base_page.py                  # 基类（自愈钩子 + 自动步骤截图）\n"
    "│   ├── login/login_page.py           # 登录页\n"
    "│   ├── home/home_page.py             # 平台首页\n"
    "│   ├── data_platform/data_platform_page.py  # 数据平台首页 + 侧边栏导航\n"
    "│   ├── monitor/overview_page.py      # 监控中心-概览页\n"
    "│   └── data_integration/source_page.py      # 数据集成-数据源字段映射页\n"
    "│\n"
    "├── common/                           # 🧱 公共组件（与 pages 同级）\n"
    "│   ├── navigation.py                 # 导航/菜单通用操作\n"
    "│   └── dialog.py                     # 通用对话框 / Toast 提示\n"
    "│\n"
    "├── utils/                            # 🔧 基础设施\n"
    "│   ├── browser.py                    # Playwright 启动 + browser/context/page 管理\n"
    "│   ├── auth.py                       # 登录态获取（UI 登录 / API Token 直取）\n"
    "│   ├── config.py                     # YAML 配置读取（支持环境变量覆盖）\n"
    "│   ├── assertions.py                 # UI 断言封装\n"
    "│   ├── screenshot.py                 # 截图命名 + Allure 附件\n"
    "│   ├── locator_parser.py             # 定位规范化工具\n"
    "│   ├── excel_export.py               # Excel 导出（PASS 绿底 / FAIL 红底）\n"
    "│   └── dingtalk.py                   # 钉钉机器人推送（签名加密）\n"
    "│\n"
    "├── tests/                            # ✅ 测试用例层\n"
    "│   ├── conftest.py                   # pytest fixtures（page / 登录态 / 失败截图）\n"
    "│   ├── test_login.py                 # 登录成功 + 密码错误用例\n"
    "│   ├── test_home_navigation.py       # 首页卡片点击 + 侧边栏导航\n"
    "│   ├── test_overview.py              # 概览页 7 个用例（7/7 通过）\n"
    "│   └── test_data_integration.py      # 数据源字段映射 4 个用例（4/4 通过）\n"
    "│\n"
    "├── scripts/\n"
    "│   ├── run_pipeline.py               # 🚀 8 阶段一键 Pipeline（核心入口）\n"
    "│   ├── run_overview_tests.py         # 🧪 概览页一键运行脚本（CI 友好）\n"
    "│   └── gen_templates.py              # 生成 Excel 模板文件\n"
    "│\n"
    "├── UI_input_files/                   # 📥 输入端（放 PRD/蓝湖/截图）\n"
    "│   ├── prd/\n"
    "│   ├── screenshots/\n"
    "│   └── lanhu/\n"
    "│\n"
    "├── UI_output_files/                  # 📤 Excel 输出\n"
    "│   ├── test_ui_results/              # PASS 绿底 / FAIL 红底\n"
    "│   └── bug_list/                     # Bug 清单（TAPD 可直接导入）\n"
    "│\n"
    "├── reports/                          # 📊 运行产物\n"
    "│   ├── allure-results/               # Allure 原始数据\n"
    "│   ├── allure-report/                # Allure HTML 报告（SPA，需 HTTP 访问）\n"
    "│   ├── screenshots/                  # 失败截图 + 步骤截图\n"
    "│   ├── self_healing/                 # 自愈定位记录\n"
    "│   └── videos/                       # 失败用例视频（如开启）\n"
    "│\n"
    "├── templates/                        # Excel 模板（bug + 测试结果）\n"
    "├── .github/workflows/test.yml        # ⚙️  GitHub Actions 配置\n"
    "├── requirements.txt\n"
    "├── pytest.ini                        # pytest 公共参数\n"
    "├── .gitignore\n"
    "└── README.md"
)

doc.add_page_break()

# ============================================================
# 四、配置说明
# ============================================================
doc.add_heading("四、配置说明", level=1)
doc.add_heading("4.1 基础配置 config/config.yaml", level=2)
doc.add_paragraph("⚠️  首次使用必须修改以下 3 处：base_url、登录账号、AI 引擎模式。配置路径： config/config.yaml。")

# 分段展示每个配置块
doc.add_paragraph()
doc.add_paragraph("① 项目与测试环境", style="List Bullet")
add_code_block(
    "project_name: \"数据平台\"\n\n"
    "test_env:\n"
    "  base_url: \"https://hrplatform.tgeem.cn\"   # ⭐ 必须修改：被测系统首页 URL"
)
doc.add_paragraph()

doc.add_paragraph("② 浏览器配置", style="List Bullet")
add_simple_table(
    ["配置项", "默认值", "说明"],
    [
        ["type", "chromium", "可选 chromium / firefox / webkit，一般用 chromium"],
        ["headless", "false", "调试用 false（可见浏览器），CI 环境 true（无头）"],
        ["slow_mo", "100", "每步延迟 100ms，调试观察用；CI 执行设为 0 可提速 30%+"],
        ["viewport", "1920x1080", "视口大小，根据被测系统实际分辨率调整"],
        ["record_video", "false", "失败用例是否录制视频（占用磁盘）"],
        ["trace", "on-first-retry", "Playwright trace 策略，可设为 retain-on-failure 排查问题"],
    ],
)
doc.add_paragraph()

doc.add_paragraph("③ 登录账号配置", style="List Bullet")
add_simple_table(
    ["配置项", "说明"],
    [
        ["username", "SSO/平台登录账号（建议用专用测试账号）"],
        ["password", "登录密码（生产环境建议用环境变量 AUTH_PASSWORD 覆盖）"],
        ["login_url", "登录页路径，默认为 /login，与 base_url 拼接"],
        ["login_api", "登录接口路径，优先调 API 获取 Token 跳过 UI 登录，失败自动回退到 UI 登录"],
        ["storage_state", "登录态持久化文件，首次登录后保存，24 小时内复用，大幅提速"],
    ],
)
doc.add_paragraph()

doc.add_paragraph("④ 超时与重试", style="List Bullet")
add_code_block(
    "timeout:\n"
    "  page_load: 30          # 页面加载超时（秒）\n"
    "  element: 10            # 元素查找超时（秒，页面元素多或网络差可加大）\n"
    "  navigation: 30         # 页面跳转超时\n"
    "\n"
    "retry:\n"
    "  max_retries: 3         # 用例失败重试次数（偶发网络抖动）\n"
    "  retry_delay: 1         # 重试间隔（秒）"
)
doc.add_paragraph()

doc.add_paragraph("⑤ AI 引擎配置（⭐ 重点）", style="List Bullet")
add_simple_table(
    ["配置项", "默认值", "说明"],
    [
        ["enabled", "true", "总开关，关掉后 Pipeline 跳过 AI 生成，直接跑已有用例"],
        ["self_healing", "true", "定位失败是否触发自愈（Midscene 风格派生候选 + LLM 兜底）"],
        ["mode", "trae", "trae：对话接管，生成 prompt 到 ai/.pending/，零配置\napi：HTTP 直调 LLM，需填下面的 Key，适合 CI 无人值守"],
        ["provider", "openai", "openai / azure / 自建（OpenAI 兼容接口即可）"],
        ["openai_base_url", "\"\"", "自建网关或代理地址，如 https://api.openai.com/v1"],
        ["api_key", "\"\"", "API Key，推荐使用环境变量 AI_API_KEY 覆盖（更安全）"],
        ["model", "gpt-4o", "识别截图时模型必须支持 vision（如 gpt-4o / claude-3-5-sonnet / qwen-vl）"],
        ["temperature", "0.2", "代码生成用低温度，更稳定"],
    ],
)
doc.add_paragraph()

doc.add_paragraph("⑥ 输入端配置", style="List Bullet")
add_code_block(
    "input:\n"
    "  mode: \"mixed\"         # prd（只解析 PRD）| lanhu | screenshot | mixed（混合交叉印证）\n"
    "  prd_dir: \"UI_input_files/prd\"\n"
    "  screenshot_dir: \"UI_input_files/screenshots\"\n"
    "  lanhu_dir: \"UI_input_files/lanhu\"\n"
    "  lanhu_base_url: \"\"    # 蓝湖分享链接，留空则读本地导出文件"
)
doc.add_paragraph()

doc.add_paragraph("⑦ 截图配置", style="List Bullet")
add_simple_table(
    ["配置项", "默认值", "说明"],
    [
        ["on_failure", "true", "用例失败自动截图（推荐保持）"],
        ["on_step", "true", "每个 allure.step 自动截图，生成可视化报告（占用磁盘，调试/回归可开启）"],
        ["full_page", "true", "是否整页滚动截图，false 则只截可视区"],
        ["dir", "reports/screenshots", "截图保存目录"],
    ],
)

doc.add_heading("4.2 钉钉通知 config/notification.json", level=2)
add_note("配置机器人时记得把钉钉「机器人安全设置」中的「自定义关键词」填一个框架消息里固定带的词，本框架默认使用「测试」作为关键词。")
add_code_block(
    "{\n"
    "  \"dingtalk\": {\n"
    "    \"webhook_url\": \"https://oapi.dingtalk.com/robot/send?access_token=【您的Token】\",\n"
    "    \"secret\": \"SECxxxxxxxxxxxxxxxx\",  # 启用加签时填写，无加签可置 \"\"\n"
    "    \"mention_mobiles\": [\"139******00\"],\n"
    "    \"at_all\": false\n"
    "  },\n"
    "  \"threshold\": {\n"
    "    \"pass_rate_warning\": 95,   # 通过率低于此值，钉钉消息标红\n"
    "    \"p0_fail_immediate\": true  # P0 失败立即 @ 指定人\n"
    "  },\n"
    "  \"project_name\": \"数据平台\",\n"
    "  \"env\": \"测试环境\"\n"
    "}"
)

doc.add_heading("4.3 Pytest 配置 pytest.ini", level=2)
add_code_block(
    "[pytest]\n"
    "addopts = -v --tb=short --alluredir=reports/allure-results\n"
    "testpaths = tests\n"
    "python_files = test_*.py\n"
    "markers =\n"
    "    P0: 核心冒烟用例\n"
    "    P1: 重要功能用例\n"
    "    P2: 边界/异常用例\n"
    "    smoke: 冒烟测试\n"
    "    regression: 回归测试\n"
    "    overview: 监控中心概览页专项测试"
)

doc.add_page_break()

# ============================================================
# 五、Pipeline 运行流程
# ============================================================
doc.add_heading("五、Pipeline 运行流程", level=1)

doc.add_heading("5.1 八阶段全链路", level=2)
doc.add_paragraph("完全对齐 API 自动化框架的流水线设计，8 个阶段依次执行：")
add_simple_table(
    ["阶段", "名称", "核心动作", "产物"],
    [
        ["[1]", "导入文档", "扫描 UI_input_files/ 下的 PRD、蓝湖、截图", "输入文件清单"],
        ["[2]", "文档解析", "parsers/ 将各类型输入解析为统一的 UIRequirement JSON", "UI_input_files/parsed_requirement.json"],
        ["[3]", "AI 生成用例", "ai/generators/ 读取 parsed_requirement.json 生成 Page Object + pytest 用例", "pages/**/*.py、tests/test_*.py"],
        ["[4]", "执行用例", "pytest -n 2 + Playwright 执行测试，含自愈与失败重试", "reports/allure-results/"],
        ["[5]", "导出结果", "将每条用例的通过/失败写入 Excel，PASS 绿色、FAIL 红色", "UI_output_files/test_ui_results/"],
        ["[5b]", "导出 Bug", "失败用例按 TAPD 模板生成 Bug 清单", "UI_output_files/bug_list/"],
        ["[6]", "生成报告", "调用 Allure generate 生成 HTML 报告，并启动本地 HTTP 服务", "reports/allure-report/index.html + http://localhost:8088"],
        ["[7]", "推送钉钉", "通过通过率 + 失败 Top 5 + 报告链接推送机器人消息", "钉钉群消息"],
        ["[8]", "CI/CD", "校验 GitHub Actions 配置并打印结论（实际 push 触发）", "通过/失败 exit code"],
    ],
)

doc.add_heading("5.2 AI 引擎详解", level=2)
doc.add_paragraph("Pipeline 的「阶段 3」是 UI 框架与传统测试框架的本质区别，支持两种 AI 运行模式：")
add_simple_table(
    ["模式", "配置 mode", "使用场景", "操作方式"],
    [
        ["trae 对话模式", "trae（默认）", "在 Trae IDE 对话框内编码，无需配置 Key",
         "生成器将 prompt 写入 ai/.pending/*.txt，Trae 对话自动识别并产出代码到 pages/ 和 tests/，完成后移至 ai/.done/"],
        ["API 直调模式", "api", "CI 环境无人值守、批量生成",
         "配置 openai_base_url + api_key + model（支持 vision 才能识别截图），生成器 HTTP 直调返回 Python 代码"],
    ],
)

doc.add_heading("5.3 输入端解析", level=2)
doc.add_paragraph("config.yaml input.mode = mixed 时，同名页面会自动合并：PRD 的文字描述 + 蓝湖的坐标信息 + 截图视觉证据。")
add_simple_table(
    ["方式", "输入", "目录", "解析逻辑"],
    [
        ["A PRD", ".docx / .pdf / .md", "UI_input_files/prd/", "按标题层级切分页面，提取元素、按钮、输入框、预期校验点"],
        ["B 蓝湖", ".json / .html 导出 或 在线链接", "UI_input_files/lanhu/", "解析图层信息，精确到像素级坐标，生成定位 hint"],
        ["C 原型截图", ".png / .jpg", "UI_input_files/screenshots/", "每张截图视为一页，AI 需使用 vision 模型识别按钮/表格/输入框"],
        ["D 混合", "任意组合", "同上", "同名页面交叉印证，优先取 PRD 文本 + 蓝湖坐标，缺失部分由截图视觉补充"],
    ],
)

doc.add_heading("5.4 自愈机制", level=2)
doc.add_paragraph("开启 config.ai.self_healing = true 后，pages/base_page.py 的 click / fill / wait_for 失败时自动进入自愈流程：")
items = [
    "语义候选派生：基于原定位器的 role / text / aria-label / placeholder / name 属性生成 10+ 候选",
    "逐个试探命中：按概率高低逐个尝试，命中则记录为替代定位，继续执行",
    "LLM 兜底：10 个候选都未命中，则调用 AI 引擎基于页面 DOM 快照重新生成定位",
    "自愈记录：结果写入 reports/self_healing/，在 Allure 附件中以 link 形式展示，便于回归时替换失效定位",
]
for item in items:
    doc.add_paragraph(item, style="List Number")
add_note("自愈命中的定位器不会自动改回源码，建议每周检查一次自愈记录，将高频命中项人工回填到 pages/**/ 对应 Page 类中。")

doc.add_page_break()

# ============================================================
# 六、使用指南
# ============================================================
doc.add_heading("六、使用指南", level=1)

doc.add_heading("6.1 快速开始（5 步跑通冒烟）", level=2)
add_code_block(
    "# Step 1. 修改 config/config.yaml\n"
    "#    base_url  → 被测环境地址\n"
    "#    username / password  → 测试账号\n"
    "#    ai.mode  → 没有 API Key 就保持 trae\n"
    "#    ai.enabled  → 首次跑建议 false（直接用已有用例）\n\n"
    "# Step 2. 安装依赖\n"
    "pip install -r requirements.txt\n"
    "python -m playwright install chromium --with-deps\n\n"
    "# Step 3. 生成 Excel 模板（首次使用）\n"
    "python scripts/gen_templates.py\n\n"
    "# Step 4. 执行一次冒烟（跳过 AI 生成，直接跑所有已存在用例）\n"
    "python scripts/run_pipeline.py --skip-gen\n\n"
    "# Step 5. 查看产物\n"
    "#   Allure 报告: reports/allure-report/index.html（已自动在 8088 端口托管）\n"
    "#   Excel 结果: UI_output_files/test_ui_results/\n"
    "#   Bug 清单:   UI_output_files/bug_list/\n"
    "#   钉钉消息:   群机器人推送（如配置 notification.json）"
)

doc.add_heading("6.2 全链路运行（含 AI 生成）", level=2)
doc.add_paragraph("当你有新的 PRD / 蓝湖 / 截图需要生成用例时，按如下流程：")
add_code_block(
    "# Step 1. 把输入文件放入对应目录\n"
    "UI_input_files/\n"
    "├── prd/                 # 放 .docx / .pdf / .md 需求文档\n"
    "├── lanhu/               # 放蓝湖导出 .json / .html\n"
    "└── screenshots/         # 放原型截图 .png / .jpg\n\n"
    "# Step 2. 修改 config.yaml\n"
    "#   ai.enabled: true\n"
    "#   ai.mode: \"trae\" 或 \"api\"\n"
    "#   input.mode: \"mixed\"（有什么就填什么）\n\n"
    "# Step 3. 只跑 AI 生成（验证产出代码不直接执行）\n"
    "python scripts/run_pipeline.py --gen-only\n"
    "# 检查 pages/ 与 tests/ 下 AI 生成的文件，必要时人工修正\n\n"
    "# Step 4. 调试单个生成的用例\n"
    "pytest tests/test_XXXX.py -v -s --tb=long  # -s 可见 print 输出\n\n"
    "# Step 5. 全链路执行\n"
    "python scripts/run_pipeline.py"
)

doc.add_heading("6.3 单模块 / 筛选运行", level=2)
doc.add_paragraph("日常开发调试不需要每次跑全量，pytest 提供多种筛选方式：")
add_code_block(
    "# 1. 按文件（单模块）\n"
    "pytest tests/test_overview.py -v\n\n"
    "# 2. 按 marker（优先级）\n"
    "pytest tests/ -m P0 -v         # 只跑 P0 核心冒烟\n"
    "pytest tests/ -m \"P0 or P1\" -v\n\n"
    "# 3. 按自定义 marker（本次新增的概览页专项）\n"
    "pytest tests/ -m overview -v\n\n"
    "# 4. 按关键字（-k 用例名/类名）\n"
    "pytest tests/ -k \"login or overview\" -v\n\n"
    "# 5. 并行（xdist）\n"
    "pytest tests/ -n 2 -v           # 2 个 worker，Windows 推荐不超过 2\n\n"
    "# 6. 失败重试 + trace\n"
    "pytest tests/test_login.py -v --reruns=2 --trace retain-on-failure\n\n"
    "# 7. 脚本方式（CI 友好）：概览页专项一键脚本\n"
    "python scripts/run_overview_tests.py --ci -n 2 --no-report"
)

doc.add_heading("6.4 AI 模式说明", level=2)
doc.add_paragraph("trae 模式 vs api 模式的操作差异：")
add_simple_table(
    ["项目", "trae 对话模式", "api HTTP 直调模式"],
    [
        ["配置", "无需配置 API Key，ai.mode=\"trae\" 即可", "需要填 openai_base_url + api_key + model"],
        ["流程", "ai/generators 把 prompt 写入 ai/.pending/，对话 AI 读取生成代码", "ai/generators 直接 HTTP 调 OpenAI 接口获取代码"],
        ["代码质量控制", "可人工在对话中迭代追问，质量更高", "一次生成，需后处理兜底"],
        ["适用场景", "日常开发、Trae IDE 内用", "CI 无人值守、批量生成回归"],
        ["产出落地位置", "pages/** 与 tests/，完成后 ai/.pending/ 移至 ai/.done/", "pages/** 与 tests/，成功即写入"],
    ],
)

doc.add_page_break()

# ============================================================
# 七、产物说明
# ============================================================
doc.add_heading("七、产物说明", level=1)

doc.add_heading("7.1 Allure HTML 报告", level=2)
doc.add_paragraph("Pipeline stage6 自动生成，目录：reports/allure-report/")
doc.add_paragraph("⭐  注意事项：")
items = [
    "Allure 是 SPA，直接双击 index.html 会一直卡在 Loading，必须用 HTTP 服务器托管。",
    "脚本已自动在 stage6 启动 http.server 监听 8088 端口：http://localhost:8088",
    "手动启动命令： python -m http.server 8088 --directory reports/allure-report",
    "报告主要标签页： Overview（总览）、Categories（缺陷分类）、Suites（用例树）、Behaviors（特性/故事）、Graphs（图表）。",
    "每个用例的 Steps 标签页可看到每个 allure.step 的截图和耗时。",
]
for item in items:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("7.2 Excel 结果", level=2)
add_simple_table(
    ["目录", "文件命名", "说明"],
    [
        ["UI_output_files/test_ui_results/", "test_ui_results_YYYYMMDD_HHMMSS.xlsx",
         "每一行一条用例：模块、用例名、优先级、P0/P1/P2、结果（PASS 绿底/FAIL 红底）、耗时、备注"],
    ],
)

doc.add_heading("7.3 Bug 清单", level=2)
add_simple_table(
    ["目录", "文件命名", "说明"],
    [
        ["UI_output_files/bug_list/", "bug_list_YYYYMMDD_HHMMSS.xlsx",
         "每条失败用例一行，TAPD 模板格式可直接导入。Bug 标题格式：【{模块}】{用例名} UI → {错误码} {错误消息}。"],
    ],
)

doc.add_heading("7.4 其他产物", level=2)
add_simple_table(
    ["目录", "说明"],
    [
        ["reports/screenshots/", "步骤级截图 + 失败截图，按文件名可关联到具体用例步骤"],
        ["reports/self_healing/", "自愈记录 JSON，记录原定位 → 命中定位 → 命中概率"],
        ["reports/videos/", "失败用例视频回放（需开启 config.browser.record_video）"],
        ["reports/allure-results/", "Allure 原始 JSON，供 allure generate 二次加工"],
    ],
)

doc.add_page_break()

# ============================================================
# 八、CI/CD
# ============================================================
doc.add_heading("八、CI/CD 集成", level=1)
doc.add_paragraph("仓库已内置 GitHub Actions 工作流：.github/workflows/test.yml，拆分了两个 Job：")

add_simple_table(
    ["Job 名", "触发条件", "运行内容", "耗时"],
    [
        ["overview-test", "每次 push 到 main / develop，或 PR 到 main",
         "python scripts/run_overview_tests.py --ci -n 2：快速跑 7 个概览页用例，验证环境可用性", "~1 分钟"],
        ["full-pipeline", "手动 workflow_dispatch，且 suite = full",
         "python scripts/run_pipeline.py：8 阶段全链路，Allure 报告 + Excel 作为 Artifact 上传", "~10 分钟"],
    ],
)

doc.add_paragraph("GitHub Repository 建议配置以下 Secrets：")
add_simple_table(
    ["Secret 名", "用途"],
    [
        ["TEST_AUTH_USERNAME", "覆盖 config.auth.username（避免密码硬编码到仓库）"],
        ["TEST_AUTH_PASSWORD", "覆盖 config.auth.password"],
        ["TEST_AUTH_TOKEN", "覆盖 login_api 方式直取 Token 的凭证（可选）"],
        ["DINGTALK_WEBHOOK", "覆盖 notification.json 中的 webhook_url（可选）"],
        ["DINGTALK_SECRET", "覆盖 notification.json 中的 secret（可选）"],
    ],
)
doc.add_paragraph()
add_code_block(
    "# 本地手动触发等价命令（模拟 CI）：\n"
    "# 概览页快速验证\n"
    "python scripts/run_overview_tests.py --ci --no-report\n\n"
    "# 全量 Pipeline\n"
    "python scripts/run_pipeline.py"
)

doc.add_page_break()

# ============================================================
# 九、常见问题 FAQ
# ============================================================
doc.add_heading("九、常见问题 FAQ", level=1)

faqs = [
    ("Q1: allure 命令找不到 / WinError 2",
     "答：Windows 下需要手动将 allure-2.34.1/bin 加入 PATH 环境变量。cmd 里执行 where allure，能看到 allure.bat 即可。"
     "修改 PATH 后务必重启 IDE 或 Shell，旧的进程池读不到新变量。"),
    ("Q2: Allure generate 报错 ClassNotFoundException io.qameta.allure.CommandLine",
     "答：原因是使用 shell=True 调用 allure.bat 时 Java classpath 丢失。正确做法：用 shutil.which(\"allure.bat\") 拿到完整路径后，以 [path, generate, ...] 列表方式调用，不要拼字符串。run_pipeline.py 中已修复。"),
    ("Q3: data_integration 用例 Locator.wait_for('table tr') 超时",
     "答：DolphinScheduler 真实页面使用 Naive UI 的 .n-data-table，且进入需要先点数据集成父菜单再点子菜单。新版本 source_page.py 已改为 .n-data-table-tbody tr 并实现两步菜单导航。"),
    ("Q4: pytest-xdist 并发出现 worker 崩溃、用例被 skip",
     "答：Windows + Python 3.14 + Playwright 多进程下 asyncio 事件循环不稳定。已将 stage4 并发数固定为 -n 2，不要用 -n auto，默认 auto 会开 16 个 worker 导致崩溃。"),
    ("Q5: 钉钉机器人返回 errcode 300005 token is not exist",
     "答：webhook_url 嵌套了两次 access_token= 前缀。正确格式是 https://oapi.dingtalk.com/robot/send?access_token=【64位Token】。notification.json 已修复。"),
    ("Q6: 钉钉推送 errcode 310000 关键词不匹配",
     "答：钉钉机器人安全设置启用了「自定义关键词」，而消息正文必须至少包含一个关键词。utils/dingtalk.py 在末尾固定添加「测试」关键词，请在机器人后台加上这个关键词，或修改 send_report() 的模板。"),
    ("Q7: 登录态 token 过期，每个用例都重新登录很慢",
     "答：修改 config.auth.storage_state 指向 config/.auth_state.json，BrowserManager 会在首次 UI 登录成功后保存。"
     "只要 SESSION 未过期（通常 24 小时），后续用例直接复用，可提速 3~5 倍。"),
    ("Q8: 概览页 Allure 报告中截图为空白或元素错位",
     "答：Playwright 默认等待 networkidle，但 ECharts 等 JS 渲染会在网络空闲后继续，可在关键步骤后显式 page.wait_for_timeout(500) 或等具体 canvas 可见。"),
    ("Q9: AI 生成的 Page Object 或 test_ 脚本是纯中文文本（无效 Python）",
     "答：这是 trae 对话模式的正常提示：ai/.pending/ 下的 prompt 还没被对话 AI 处理。"
     "请等待对话 AI 读取并产出代码，或切换 ai.mode: api + 配置 API Key 走 HTTP 直调。run_pipeline.py stage3 已过滤中文占位文本不再写入。"),
    ("Q10: Pipeline 明明有失败用例，stage8 却显示『Pipeline 通过』",
     "答：旧版本 stage8 只判断 failed==0，忽略了 broken。修复后 collect_stats 按 (name|fullName) 去重 rerun 记录，且 stage8 通过条件改为 failed==0 and broken==0。已在最新 commit 中修复。"),
    ("Q11: 如何迁移到新项目？",
     "答：修改 4 处即可复用：\n"
     "  1) config/config.yaml: base_url / 账号 / 登录路径 / AI 模式\n"
     "  2) config/notification.json: 钉钉机器人\n"
     "  3) 新项目 PRD / 蓝湖 / 截图放入 UI_input_files/\n"
     "  4) pages/ 与 tests/：AI 生成或手写业务页\n"
     "  框架核心 ai/、parsers/、utils/、conftest.py、run_pipeline.py 可 100% 复用。"),
]

for q, a in faqs:
    pq = doc.add_paragraph()
    rq = pq.add_run(q)
    rq.font.bold = True
    rq.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    pa = doc.add_paragraph()
    ra = pa.add_run(a)
    ra.font.size = Pt(10.5)

# ============================================================
# 保存
# ============================================================
doc.save(OUT)
print(f"✅ 手册已生成: {OUT}  (size={OUT.stat().st_size} bytes)")
